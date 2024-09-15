import os
from typing import Dict
from PyPDF2 import PdfReader
from docx import Document
from uuid import uuid4

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents.base import Document


from src.app.core.logs import logger
from src.app.utils import create_qdrant_collection, init_vector_store

def extract_text_from_pdf(file_path):
    try:
        pdf = PdfReader(open(file_path, 'rb'))
        text = ''
        num_pages = len(pdf.pages)
        for page_num in range(num_pages):
            text += pdf.pages[page_num].extract_text()
        return text
    except Exception as e:
        logger.info(f"An error occurred while reading PDF: {e}")
        return None


def extract_text_from_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.info(f"An error occurred while reading TXT: {e}")
        return None


def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        text = ''
        for para in doc.paragraphs:
            text += para.text + '\n'
        return text
    except Exception as e:
        logger.info(f"An error occurred while reading DOCX: {e}")
        return None


def split_docs_docx(texts):
    """Split DocX docs into chunks.

    Args:
        texts (Iterable[Document]): DocX document with metadata.

    Returns:
        List[Document]: Splitted DocX Document

    """
    character_splitter = RecursiveCharacterTextSplitter(
        chunk_size=5000,
        chunk_overlap=100,
        is_separator_regex=False,
    )

    list_documents = character_splitter.split_text(texts)
    return list_documents, character_splitter


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from files of type PDF, TXT, and DOCX.

    params:
    -------
        file_path: str
            The path to the file.

    returns:    
    --------
        text: str
            The extracted text.
    """

    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.txt'):
        return extract_text_from_txt(file_path)
    elif file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    else:
        return None


def extract_metadata(file_path: str) -> Dict[str, str]:
    """
    Extract all possible metadata from the file,
    such as file name, size, extension, etc.

    params: 
    -------
        file_path: str
            The path to the file.

    returns:    
    --------
        metadata: dict
            A dictionary containing all the metadata.   

    """
    try:
        return {
            "file_name": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "file_extension": os.path.splitext(file_path)[1],
            "file_type": os.path.splitext(file_path)[1].replace(".", ""),
        }
    except Exception as e:
        logger.info(f"An error occurred while extracting metadata: {e}")
        return {}


def text_chunking_and_qdrant_upload(text: str, file_metadata: Dict):
    """
    Chunk the text into smaller chunks and upload them to Qdrant.

    params: 
    -------
        text: str
            The text to be chunked.
        file_metadata: dict
            A dictionary containing the metadata of the file.
        userId: str
            The user id of the user who uploaded the file.

    """
    create_qdrant_collection()
    vector_store = init_vector_store()
    list_text, character_splitter = split_docs_docx(text)
    
    logger.info(
        "Started Chunking text and extracting metadata from provided content.")

    
    metadata = {
        "file_name": file_metadata["file_name"],
        "file_size": file_metadata["file_size"],
        "file_type": file_metadata["file_type"],
        "chat_id" : file_metadata["chat_id"],
        "user_id" : file_metadata["user_id"],
    }

    langchain_documents = [
        Document(page_content=text, metadata=metadata) for text in list_text
    ]
    

    langchain_documents = character_splitter.split_documents(langchain_documents)
    langchain_documents_uuid = [str(uuid4()) for _ in range(len(langchain_documents))]

    for doc in langchain_documents:
        doc.metadata.update(
            {
                "chunk_length": len(doc.page_content)
            }
        )
    
    vector_store.add_documents(documents=langchain_documents, ids=langchain_documents_uuid)
