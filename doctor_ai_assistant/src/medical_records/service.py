import asyncio
import boto3
import uuid
import hashlib
import tiktoken
from typing import List, Dict, Any, Optional
from datetime import datetime
from pathlib import Path

from src.medical_records.models import DocumentType
from src.qdrant.client import qdrant_manager
from src.qdrant.collections import collection_manager
from src.qdrant.models import ChunkMetadata, ChunkPayload
from src.config import settings
from src.shared.logs import logger
from src.chat.constants import ModelsEnum
from qdrant_client.http import models
from langchain_openai import OpenAIEmbeddings



class MedicalRecordsFormatter:
    """Formatter for medical records data and responses"""
    
    @staticmethod
    def format_file_size(size_bytes: int) -> str:
        """Format file size in human readable format"""
        if size_bytes == 0:
            return "0 B"
        
        size_names = ["B", "KB", "MB", "GB"]
        i = 0
        while size_bytes >= 1024 and i < len(size_names) - 1:
            size_bytes /= 1024.0
            i += 1
        
        return f"{size_bytes:.1f} {size_names[i]}"
    
    @staticmethod
    def format_ingestion_summary(result: Dict[str, Any]) -> str:
        """Format ingestion result into human readable summary"""
        total = result.get("total_files", 0)
        successful = result.get("successful_ingestions", 0)
        failed = result.get("failed_ingestions", 0)
        skipped = result.get("skipped_ingestions", 0)
        chunks = result.get("total_chunks_created", 0)
        
        summary = f"Ingestion completed: {successful}/{total} files successful"
        
        if failed > 0:
            summary += f", {failed} failed"
        if skipped > 0:
            summary += f", {skipped} skipped"
        
        summary += f". Created {chunks} searchable chunks."
        
        return summary
    
    @staticmethod
    def format_search_results(results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Format search results for better presentation"""
        formatted_results = []
        
        for result in results:
            formatted_result = {
                "relevance_score": f"{result.get('score', 0):.3f}",
                "document_type": result.get('document_type', 'unknown').replace('_', ' ').title(),
                "file_name": result.get('file_name', 'unknown'),
                "content_preview": result.get('content', '')[:200] + "..." if len(result.get('content', '')) > 200 else result.get('content', ''),
                "s3_location": result.get('s3_path', ''),
                "metadata": result.get('metadata', {}),
                "chunk_id": result.get('chunk_id', '')
            }
            formatted_results.append(formatted_result)
        
        return formatted_results



class S3MedicalRecordsService:
    """Service for ingesting medical records from S3 into Qdrant"""
    
    def __init__(self):
        self.s3_client = boto3.client(
            's3',
            aws_access_key_id=settings.aws_access_key_id,
            aws_secret_access_key=settings.aws_secret_access_key,
            region_name=settings.aws_region
        )
        
        self.embedding_service = OpenAIEmbeddings(
            openai_api_key=settings.openai_api_key,
            model="text-embedding-3-small"
        )
        
        self.encoding = tiktoken.encoding_for_model(ModelsEnum.OPENAI_GPT.value)
        
        self.chunk_size = 800
        self.chunk_overlap = 100
        self.max_concurrent = 5

    async def download_s3_file(self, s3_key: str, bucket_name: str = None) -> Optional[bytes]:
        """Download file content from S3"""
        try:            
            response = self.s3_client.get_object(Bucket=bucket_name, Key=s3_key)
            content = response['Body'].read()
            
            return content
            
        except Exception as e:
            logger.exception(f"Failed to download S3 file {s3_key}: {e}")
            return None

    def classify_medical_document(self, filename: str, content: str) -> DocumentType:
        """Classify medical document type based on filename and content"""
        filename_lower = filename.lower()
        
        if any(term in filename_lower for term in ['lab', 'blood', 'urine', 'test']):
            return DocumentType.LAB_REPORT
        elif any(term in filename_lower for term in ['xray', 'ct', 'mri', 'ultrasound', 'imaging']):
            return DocumentType.IMAGING
        elif any(term in filename_lower for term in ['discharge', 'summary']):
            return DocumentType.DISCHARGE_SUMMARY
        elif any(term in filename_lower for term in ['prescription', 'medication', 'rx']):
            return DocumentType.PRESCRIPTION
        elif any(term in filename_lower for term in ['consultation', 'visit', 'note']):
            return DocumentType.CONSULTATION_NOTE
        elif any(term in filename_lower for term in ['surgery', 'surgical', 'operation']):
            return DocumentType.SURGICAL_REPORT
        else:
            return DocumentType.CLINICAL_REPORT
    
    def extract_text_from_content(self, content: bytes, filename: str) -> str:
        """Extract text content from file bytes"""
        try:
            filename_lower = filename.lower()
            
            if filename_lower.endswith('.txt'):
                return content.decode('utf-8', errors='ignore')
            elif filename_lower.endswith('.docx'):
                return self._extract_docx_text(content)
            elif filename_lower.endswith('.pdf'):
                return self._extract_pdf_text(content)
            else:
                try:
                    return content.decode('utf-8', errors='ignore')
                except:
                    return f"[Binary content from {filename}]"
        except Exception as e:
            logger.exception(f"Error extracting text from {filename}: {e}")
            return f"[Error extracting content from {filename}]"
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX file content."""
        try:
            from docx import Document
            import io
            
            doc = Document(io.BytesIO(content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.exception(f"Error extracting DOCX text: {e}")
            return "[Error extracting DOCX content]"
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF file content."""
        try:
            from PyPDF2 import PdfReader
            import io
            
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            
            for page in reader.pages:
                if page.extract_text().strip():
                    text_parts.append(page.extract_text())
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.exception(f"Error extracting PDF text: {e}")
            return "[Error extracting PDF content]"

    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """Split text into overlapping chunks using tiktoken for accurate token counting"""
        chunk_size = chunk_size or self.chunk_size
        overlap = overlap or self.chunk_overlap
        
        try:
            tokens = self.encoding.encode(text)
            
            if len(tokens) <= chunk_size:
                return [text]
            
            chunks = []
            start = 0
            
            while start < len(tokens):
                end = min(start + chunk_size, len(tokens))
                chunk_tokens = tokens[start:end]
                chunk_text = self.encoding.decode(chunk_tokens)
                if end < len(tokens):
                    last_sentences = chunk_text[-200:] if len(chunk_text) > 200 else chunk_text
                    sentence_endings = ['.', '!', '?', '\n\n']
                    
                    best_break = -1
                    for ending in sentence_endings:
                        pos = last_sentences.rfind(ending)
                        if pos > best_break:
                            best_break = pos
                    
                    if best_break > 0:
                        adjusted_end = len(chunk_text) - len(last_sentences) + best_break + 1
                        chunk_text = chunk_text[:adjusted_end]
                        actual_tokens = self.encoding.encode(chunk_text)
                        end = start + len(actual_tokens)
                
                chunks.append(chunk_text.strip())
                start = end - overlap
                
                if start <= 0 or start >= len(tokens):
                    break
            
            return [chunk for chunk in chunks if chunk.strip()]
            
        except Exception as e:
            logger.error(f"Error in tiktoken chunking: {e}")
            return self._fallback_chunk_text(text, chunk_size * 4, overlap * 4)
    
    def _fallback_chunk_text(self, text: str, char_chunk_size: int, char_overlap: int) -> List[str]:
        """Fallback character-based chunking if tiktoken fails"""
        if len(text) <= char_chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + char_chunk_size
            chunk = text[start:end]
            
            if end < len(text):
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > start + char_chunk_size // 2:
                    chunk = text[start:start + break_point + 1]
                    end = start + break_point + 1
            
            chunks.append(chunk.strip())
            start = end - char_overlap
            
            if start >= len(text):
                break
        
        return [chunk for chunk in chunks if chunk.strip()]

    async def _ingest_single_file(
        self,
        s3_key: str,
        patient_id: str,
        doctor_id: str,
        bucket_name: str = None
    ) -> Dict[str, Any]:
        """Ingest a single file from S3 with comprehensive metadata extraction."""
        try:
            content_bytes = await self.download_s3_file(s3_key, bucket_name)
            if not content_bytes:
                raise ValueError(f"Failed to download file from S3: {s3_key}")
            
            filename = Path(s3_key).name
            text_content = self.extract_text_from_content(content_bytes, filename)
            
            if not text_content or text_content.strip() == "":
                logger.warning(f"No text content extracted from {s3_key}")
                return {"status": "skipped", "reason": "no_text_content"}
            
            doc_type = self.classify_medical_document(filename, text_content)

            chunks = self.chunk_text(text_content)
            logger.info(f"Created {len(chunks)} chunks from {filename} using tiktoken")
            
            if not chunks:
                return {"status": "skipped", "reason": "no_chunks_created"}
            

            logger.debug(f"Generating embeddings for {len(chunks)} chunks")
            embeddings = self.embedding_service.embed_documents(chunks)
            
            await collection_manager.ensure_user_collections(patient_id)
            collection_name = f"user_{patient_id}_records"
            
            points = []
            for i, chunk_text in enumerate(chunks):
                chunk_id = f"{hashlib.md5(s3_key.encode()).hexdigest()}_{i}"
                
                chunk_metadata = ChunkMetadata(
                    doctor_id=doctor_id,
                    patient_id=patient_id,
                    chunk_id=chunk_id,
                    file_s3_path=s3_key,
                    medical_record_type=doc_type.value,
                    file_name=filename,
                    file_type=Path(filename).suffix.lower() or "unknown",
                    file_size=len(content_bytes),
                    file_content_type=self._get_content_type(filename)
                )
                
                chunk_payload = ChunkPayload(
                    chunk_metadata=chunk_metadata,
                    content=chunk_text
                )
                
                point = models.PointStruct(
                    id=str(uuid.uuid4()),
                    vector=embeddings[i],
                    payload=chunk_payload.dict()
                )
                points.append(point)
            
            qdrant_manager.client.upsert(
                collection_name=collection_name,
                points=points
            )
                        
            return {
                "status": "success",
                "s3_key": s3_key,
                "chunks_created": len(points),
                "document_type": doc_type.value,
                "file_size": len(content_bytes),
                "collection": collection_name
            }
            
        except Exception as e:
            logger.exception(f"Error ingesting S3 file {s3_key}: {e}")
            return {
                "status": "error",
                "s3_key": s3_key,
                "error": str(e)
            }

    def _get_content_type(self, filename: str) -> str:
        """Determine content type from filename"""
        ext = Path(filename).suffix.lower()
        content_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.txt': 'text/plain',
            '.rtf': 'application/rtf'
        }
        return content_types.get(ext, 'application/octet-stream')

    async def ingest_patient_records(
        self,
        patient_id: str,
        doctor_id: str,
        s3_keys: List[str],
    ) -> Dict[str, Any]:
        """
        Ingest medical records for a specific patient from S3.
        
        Args:
            patient_id: Patient ID
            doctor_id: Doctor ID who is uploading the records
            s3_keys: List of S3 keys for medical record files
            
        Returns:
            Dictionary with comprehensive ingestion results
        """
        try:
            valid_keys, invalid_keys = MedicalRecordsValidator.validate_s3_keys_batch(s3_keys)
            
            if invalid_keys:
                logger.warning(f"Invalid S3 keys found: {invalid_keys}")
            
            if not valid_keys:
                raise ValueError("No valid S3 keys provided")
                        
            semaphore = asyncio.Semaphore(self.max_concurrent)
            
            async def process_file(s3_key: str):
                async with semaphore:
                    return await self._ingest_single_file(s3_key, patient_id, doctor_id, settings.s3_bucket_name)
            
            results = await asyncio.gather(
                *[process_file(s3_key) for s3_key in valid_keys],
                return_exceptions=True
            )
            
            successful_ingestions = 0
            failed_ingestions = 0
            skipped_ingestions = 0
            ingestion_details = []
            total_chunks = 0
            
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    failed_ingestions += 1
                    ingestion_details.append({
                        "s3_key": valid_keys[i],
                        "status": "error",
                        "error": str(result)
                    })
                elif result["status"] == "success":
                    successful_ingestions += 1
                    total_chunks += result.get("chunks_created", 0)
                    ingestion_details.append(result)
                elif result["status"] == "skipped":
                    skipped_ingestions += 1
                    ingestion_details.append(result)
                else:
                    failed_ingestions += 1
                    ingestion_details.append(result)
            
            final_result = {
                "patient_id": patient_id,
                "doctor_id": doctor_id,
                "total_files": len(valid_keys),
                "invalid_files": invalid_keys,
                "successful_ingestions": successful_ingestions,
                "failed_ingestions": failed_ingestions,
                "skipped_ingestions": skipped_ingestions,
                "total_chunks_created": total_chunks,
                "ingestion_details": ingestion_details,
                "summary": MedicalRecordsFormatter.format_ingestion_summary({
                    "total_files": len(valid_keys),
                    "successful_ingestions": successful_ingestions,
                    "failed_ingestions": failed_ingestions,
                    "skipped_ingestions": skipped_ingestions,
                    "total_chunks_created": total_chunks
                }),
                "timestamp": datetime.now().isoformat()
            }
            
            return final_result
            
        except Exception as e:
            logger.exception(f"Error ingesting records for patient {patient_id}: {e}")
            raise

    async def search_patient_records(
        self,
        patient_id: str,
        query_text: str,
        doctor_id: str = None,
        limit: int = 10,
        score_threshold: float = 0.5,
        document_types: List[DocumentType] = None
    ) -> List[Dict[str, Any]]:
        """
        Search medical records for a specific patient with full metadata.
        
        Args:
            patient_id: Patient ID to search records for
            query_text: Search query text
            doctor_id: Optional doctor ID filter
            limit: Maximum number of results
            score_threshold: Minimum relevance score
            document_types: Optional list of document types to filter by
            
        Returns:
            List of search results with full metadata
        """
        try:
            query_embedding = self.embedding_model.embed_query(query_text)
            
            collection_name = f"user_{patient_id}_records"
            
            if not qdrant_manager.collection_exists(collection_name):
                return []
            
            filters = []
            
            if doctor_id:
                filters.append(
                    models.FieldCondition(
                        key="chunk_metadata.doctor_id",
                        match=models.MatchValue(value=doctor_id)
                    )
                )
            
            if document_types:
                doc_type_values = [dt.value for dt in document_types]
                filters.append(
                    models.FieldCondition(
                        key="chunk_metadata.medical_record_type",
                        match=models.MatchAny(any=doc_type_values)
                    )
                )
            
            search_filter = None
            if filters:
                search_filter = models.Filter(must=filters)
            
            search_results = qdrant_manager.client.search(
                collection_name=collection_name,
                query_vector=query_embedding,
                query_filter=search_filter,
                limit=limit,
                score_threshold=score_threshold,
                with_payload=True
            )
            
            formatted_results = []
            for result in search_results:
                formatted_result = {
                    "score": result.score,
                    "point_id": result.id,
                    "content": result.payload.get("content", ""),
                    "metadata": result.payload.get("chunk_metadata", {}),
                    "s3_path": result.payload.get("chunk_metadata", {}).get("file_s3_path", ""),
                    "document_type": result.payload.get("chunk_metadata", {}).get("medical_record_type", ""),
                    "file_name": result.payload.get("chunk_metadata", {}).get("file_name", ""),
                    "chunk_id": result.payload.get("chunk_metadata", {}).get("chunk_id", "")
                }
                formatted_results.append(formatted_result)
            
            return formatted_results
            
        except Exception as e:
            logger.exception(f"Error searching records for patient {patient_id}: {e}")
            raise

s3_medical_records_service = S3MedicalRecordsService()
